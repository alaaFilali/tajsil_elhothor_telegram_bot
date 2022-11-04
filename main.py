from telebot import TeleBot
from datetime import datetime
from docx import Document
import os
import sys

application_path = os.path.dirname(sys.executable)

API_KEY = os.getenv('API_KEY')

bot = TeleBot(API_KEY)

document = Document()

student_ids = []
prof_name = ''
begin = False


@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(message.chat.id, "Hello there! if you are a prof type /prof"
                                      " and if your are a student please wait for your prof to create a file"
                                      " and then type /reg")


@bot.message_handler(commands=["reg"])
def reg(message):
    global student_ids
    global prof_name
    chat_id = message.chat.id

    if chat_id not in student_ids and prof_name != '':
        bot.send_message(message.chat.id, 'Hi there! please send your name and your group.')
    elif prof_name == '':
        bot.send_message(message.chat.id, "Please wait for your prof to create a file and then retry.")
    else:
        bot.send_message(message.chat.id, "I'm sorry, you cannot register your self twice.")


@bot.message_handler(commands=['end'])
def end(message):
    global prof_name
    global document
    bot.send_message(message.chat.id, f'Sending the file to {prof_name}....')
    file_name = f'list_of_students-{datetime.today().strftime("%Y-%m-%d")}.docx'
    final_path = os.path.join(application_path, file_name)
    document.save(final_path)
    doc = open(final_path, 'rb')
    bot.send_document(message.chat.id, doc)
    bot.send_message(message.chat.id, "Done!.")


@bot.message_handler(commands=['prof'])
def get_prof(message):
    global begin
    begin = True
    bot.send_message(message.chat.id, "Hello professor, please send your name.")


def get_class_name(message):
    if message.text.lower().split()[0] == 'class':
        return True
    else:
        return False


@bot.message_handler(func=get_class_name)
def class_name(message):
    global prof_name
    global document
    class_ = message.text.lower().split()[1:]
    document.add_heading(f'Class: {" ".join(class_)}\nDateTime: {datetime.today().strftime("%Y-%m-%d, %H:%M")}\n'
                         f'Prof: {prof_name.capitalize()}\n\n', 1)
    bot.send_message(message.chat.id, "Class name added to the file with today's date and time.")
    bot.send_message(message.chat.id, 'Please tell your students to use the /reg keyword to register them selves'
                                      ' and when they are done use the command /end .')


def has_numbers(input_string):
    return any(char.isdigit() for char in input_string)


def get_name(message):
    global prof_name
    global begin
    if has_numbers(message.text):
        return True
    elif prof_name == '' and begin:
        prof_name = message.text.lower()
        bot.send_message(message.chat.id, 'Done.')
        bot.send_message(message.chat.id, 'Please send the class name and start by the keyword "class" .')
        return False
    elif not begin:
        bot.send_message(message.chat.id, "Hello there! if you are a prof type /prof"
                                          " and if your are a student please wait for your prof to create a file"
                                          " and then type /reg")
        return False
    else:
        bot.send_message(message.chat.id, 'Please provide your group also.')
        return False


@bot.message_handler(func=get_name)
def get_names(message):
    global document
    global prof_name
    chat_id = message.chat.id
    name = message.text
    if prof_name != '':
        bot.send_message(message.chat.id, 'Your name has been registered, thank you.')
        student_ids.append(chat_id)
        document.add_paragraph(f"{name}\n", style="List Bullet")
    else:
        bot.send_message(message.chat.id, 'Please wait for your prof to create a file and then retry.')


bot.polling()
